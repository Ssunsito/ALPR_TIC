from __future__ import annotations

import json
import time
from pathlib import Path
from difflib import SequenceMatcher

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


def norm_plate(stem: str) -> str | None:
    s = "".join(ch for ch in stem.upper() if ch.isalnum())
    return s if 5 <= len(s) <= 8 else None


def cer(pred: str, gt: str) -> float:
    # Aproximacion simple usando distancia de Levenshtein por DP.
    a, b = pred or "", gt or ""
    if len(b) == 0:
        return 0.0 if len(a) == 0 else 1.0
    dp = [[0] * (len(b) + 1) for _ in range(len(a) + 1)]
    for i in range(len(a) + 1):
        dp[i][0] = i
    for j in range(len(b) + 1):
        dp[0][j] = j
    for i in range(1, len(a) + 1):
        for j in range(1, len(b) + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            dp[i][j] = min(
                dp[i - 1][j] + 1,
                dp[i][j - 1] + 1,
                dp[i - 1][j - 1] + cost,
            )
    return dp[len(a)][len(b)] / max(1, len(b))


def run_image_analyzer(repo_root: Path, out_dir: Path, limit: int = 20) -> dict:
    import sys

    deploy_dir = repo_root / "deploy"
    if str(deploy_dir) not in sys.path:
        sys.path.insert(0, str(deploy_dir))

    from inference_jetson import PipelineJetson  # noqa
    from runtime_config import CONFIG, apply_runtime_profile  # noqa

    cfg = dict(CONFIG)
    apply_runtime_profile(cfg, "jetson-stable")
    cfg["show_window"] = False
    cfg["detector_model"] = str((repo_root / "models" / "plate_detector_best.pt").resolve())
    cfg["ocr_engine"] = "ctc"
    cfg["ocr_model"] = str((repo_root / "models" / "optimized" / "ocr_crnn_ctc_int8.tflite").resolve())

    pipeline = PipelineJetson(cfg)

    val_dir = repo_root / "dataset_alpr" / "images" / "val"
    ann_dir = out_dir / "anotadas"
    ann_dir.mkdir(parents=True, exist_ok=True)
    imgs = sorted(
        [
            p
            for p in val_dir.iterdir()
            if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
            and "_annotated" not in p.stem.lower()
        ]
    )
    samples = imgs[:limit]

    rows = []
    for p in samples:
        gt = norm_plate(p.stem)
        t0 = time.time()
        ann_path = ann_dir / f"{p.stem}_annotated{p.suffix}"
        _, dets = pipeline.run_single_image(str(p), output_path=str(ann_path), show_window=False)
        t1 = time.time()

        det = dets[0] if dets else {}
        pred_raw = str(det.get("text", "")).strip()
        pred = "".join(ch for ch in pred_raw.upper() if ch.isalnum())
        det_conf = float(det.get("confidence", 0.0)) if dets else 0.0
        detected = len(dets) > 0
        exact = bool(detected and gt is not None and pred == gt)
        sim = SequenceMatcher(None, pred, gt or "").ratio() if gt else 0.0
        row = {
            "image": p.name,
            "gt": gt or "",
            "pred": pred,
            "detected": int(detected),
            "exact_match": int(exact),
            "det_conf": det_conf,
            "latency_ms": (t1 - t0) * 1000.0,
            "similarity": sim,
            "cer": cer(pred, gt or "") if gt else np.nan,
        }
        rows.append(row)

    df = pd.DataFrame(rows)
    tables_dir = out_dir / "tablas"
    tables_dir.mkdir(parents=True, exist_ok=True)
    detail_csv = tables_dir / "metricas_detalle_imagen.csv"
    df.to_csv(detail_csv, index=False, encoding="utf-8")

    total = len(df)
    detected_count = int(df["detected"].sum())
    exact_count = int(df["exact_match"].sum())
    gt_count = int((df["gt"] != "").sum())

    summary = {
        "processed": total,
        "with_gt": gt_count,
        "detected_any": detected_count,
        "detection_rate": (detected_count / total) if total else 0.0,
        "exact_match": exact_count,
        "exact_match_over_gt": (exact_count / gt_count) if gt_count else 0.0,
        "latency_ms_avg": float(df["latency_ms"].mean()) if total else 0.0,
        "latency_ms_p95": float(df["latency_ms"].quantile(0.95)) if total else 0.0,
        "fps_estimated": float(1000.0 / df["latency_ms"].mean()) if total and df["latency_ms"].mean() > 0 else 0.0,
        "pred_conf_avg": float(df["det_conf"].mean()) if total else 0.0,
        "mean_similarity": float(df["similarity"].mean()) if total else 0.0,
        "mean_cer": float(df["cer"].dropna().mean()) if total else 0.0,
    }

    summary_csv = tables_dir / "metricas_resumen.csv"
    pd.DataFrame([summary]).to_csv(summary_csv, index=False, encoding="utf-8")

    summary_json = out_dir / "metricas_resumen.json"
    summary_json.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    return {
        "summary": summary,
        "detail_csv": detail_csv,
        "summary_csv": summary_csv,
        "summary_json": summary_json,
        "detail_df": df,
    }


def load_existing_metrics(repo_root: Path) -> dict:
    base = repo_root / "archive" / "temporal" / "evidencias" / "outputs"
    out = {}

    latest = base / "validation_metrics_latest.json"
    if latest.exists():
        out["validation_metrics_latest"] = json.loads(latest.read_text(encoding="utf-8"))

    history = base / "ocr_crnn_training_history.json"
    if history.exists():
        out["ocr_crnn_training_history"] = json.loads(history.read_text(encoding="utf-8"))

    yolo_csv = repo_root / "runs" / "detect" / "runs" / "plate_detector" / "yolov8n_plate_finetune" / "results.csv"
    if yolo_csv.exists():
        out["yolo_results_df"] = pd.read_csv(yolo_csv)

    return out


def generate_plots(out_dir: Path, eval_data: dict, existing: dict) -> dict:
    plots_dir = out_dir / "graficas"
    plots_dir.mkdir(parents=True, exist_ok=True)

    produced = {}

    # 1) YOLO metrics by epoch
    yolo_df = existing.get("yolo_results_df")
    if yolo_df is not None and not yolo_df.empty:
        plt.figure(figsize=(10, 6))
        plt.plot(yolo_df["epoch"], yolo_df["metrics/precision(B)"], label="Precision")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/recall(B)"], label="Recall")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50(B)"], label="mAP50")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50-95(B)"], label="mAP50-95")
        plt.xlabel("Epoca")
        plt.ylabel("Valor")
        plt.title("Evolucion de metricas YOLO por epoca")
        plt.grid(True, alpha=0.3)
        plt.legend()
        p = plots_dir / "01_yolo_metricas_epoca.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["yolo_metrics"] = p

    # 2) OCR training curves
    hist = existing.get("ocr_crnn_training_history")
    if hist:
        loss = hist.get("loss", [])
        val_loss = hist.get("val_loss", [])
        epochs = np.arange(1, len(loss) + 1)
        if len(loss) > 0 and len(val_loss) == len(loss):
            plt.figure(figsize=(10, 6))
            plt.plot(epochs, loss, label="loss")
            plt.plot(epochs, val_loss, label="val_loss")
            plt.xlabel("Epoca")
            plt.ylabel("Loss")
            plt.title("Curvas de entrenamiento OCR CRNN")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "02_ocr_loss_vs_valloss.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["ocr_loss"] = p

        blank = hist.get("val_blank_rate", [])
        empty = hist.get("val_empty_pred_rate", [])
        if len(blank) == len(empty) and len(blank) > 0:
            plt.figure(figsize=(10, 6))
            plt.plot(np.arange(1, len(blank) + 1), blank, label="val_blank_rate")
            plt.plot(np.arange(1, len(empty) + 1), empty, label="val_empty_pred_rate")
            plt.xlabel("Epoca")
            plt.ylabel("Tasa")
            plt.title("Control de colapso CTC en validacion")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "03_ocr_blank_empty_rate.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["ocr_collapse"] = p

    # 3) Bar chart old vs new summary
    old_summary = (existing.get("validation_metrics_latest") or {}).get("summary", {})
    new_summary = eval_data.get("summary", {})
    keys = ["detection_rate", "exact_match_over_gt", "mean_detector_confidence", "mean_similarity_on_detected"]
    labels = ["Detection rate", "Exact match", "Detector conf", "Similarity"]

    old_vals = [
        float(old_summary.get("detection_rate", np.nan)),
        float(old_summary.get("ocr_exact_match_overall", np.nan)),
        float(old_summary.get("mean_detector_confidence", np.nan)),
        float(old_summary.get("mean_similarity_on_detected", np.nan)),
    ]
    new_vals = [
        float(new_summary.get("detection_rate", np.nan)),
        float(new_summary.get("exact_match_over_gt", np.nan)),
        float(new_summary.get("pred_conf_avg", np.nan)),
        float(new_summary.get("mean_similarity", np.nan)),
    ]

    if not all(np.isnan(v) for v in old_vals) and not all(np.isnan(v) for v in new_vals):
        x = np.arange(len(labels))
        w = 0.35
        plt.figure(figsize=(10, 6))
        plt.bar(x - w / 2, old_vals, width=w, label="Historico")
        plt.bar(x + w / 2, new_vals, width=w, label="Analizador imagen")
        plt.xticks(x, labels)
        plt.ylim(0, 1.05)
        plt.ylabel("Valor")
        plt.title("Comparativo de metricas clave")
        plt.grid(True, axis="y", alpha=0.3)
        plt.legend()
        p = plots_dir / "04_comparativo_historico_vs_actual.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["comparison"] = p

    # 4) Latency distribution
    df = eval_data.get("detail_df")
    if df is not None and not df.empty:
        plt.figure(figsize=(10, 6))
        plt.hist(df["latency_ms"], bins=10)
        plt.xlabel("Latencia por imagen (ms)")
        plt.ylabel("Frecuencia")
        plt.title("Distribucion de latencia del analizador de imagenes")
        plt.grid(True, alpha=0.3)
        p = plots_dir / "05_histograma_latencia.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["latency_hist"] = p

        plt.figure(figsize=(10, 6))
        plt.scatter(df["det_conf"], df["similarity"], alpha=0.8)
        plt.xlabel("Confianza detector")
        plt.ylabel("Similitud OCR vs GT")
        plt.title("Relacion confianza detector vs calidad OCR")
        plt.grid(True, alpha=0.3)
        p = plots_dir / "06_scatter_confianza_vs_similitud.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["conf_vs_similarity"] = p

    return produced


def build_docx(out_dir: Path, eval_data: dict, existing: dict, plots: dict) -> Path:
    doc = Document()
    doc.add_heading("Informe de Resultados - Deteccion y OCR de Placas", 0)
    doc.add_paragraph(
        "Este documento resume las metricas cuantitativas del proyecto usando: "
        "(1) historicos de entrenamiento/validacion y (2) una corrida actual del analizador de imagenes."
    )

    summary = eval_data["summary"]
    doc.add_heading("1. Resumen de metricas (corrida actual)", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Imagenes evaluadas: {summary['processed']}\n")
    p.add_run(f"Detection rate: {summary['detection_rate']:.3f}\n")
    p.add_run(f"Exact match OCR: {summary['exact_match_over_gt']:.3f}\n")
    p.add_run(f"Latencia promedio: {summary['latency_ms_avg']:.1f} ms\n")
    p.add_run(f"Latencia p95: {summary['latency_ms_p95']:.1f} ms\n")
    p.add_run(f"FPS estimado: {summary['fps_estimated']:.3f}\n")
    p.add_run(f"Confianza promedio detector: {summary['pred_conf_avg']:.3f}\n")
    p.add_run(f"Similitud promedio OCR: {summary['mean_similarity']:.3f}\n")
    p.add_run(f"CER promedio: {summary['mean_cer']:.3f}")

    doc.add_heading("2. Graficas y analisis", level=1)

    explanations = [
        (
            "01_yolo_metricas_epoca.png",
            "Evolucion de metricas YOLO",
            "Muestra la convergencia del detector. Debe observarse estabilidad en mAP50/mAP50-95 y una brecha controlada entre precision y recall.",
        ),
        (
            "02_ocr_loss_vs_valloss.png",
            "Curvas de entrenamiento OCR",
            "Permite verificar aprendizaje y posible sobreajuste. Si val_loss deja de bajar mientras loss sigue bajando, hay indicios de sobreajuste.",
        ),
        (
            "03_ocr_blank_empty_rate.png",
            "Tasas de blank y prediccion vacia",
            "Estas curvas evalúan colapso CTC. Una reduccion y estabilizacion indica que el OCR evita producir secuencias vacias.",
        ),
        (
            "04_comparativo_historico_vs_actual.png",
            "Comparativo historico vs corrida actual",
            "Compara deteccion, exactitud OCR, confianza y similitud para evidenciar mejora o degradacion respecto a validaciones previas.",
        ),
        (
            "05_histograma_latencia.png",
            "Distribucion de latencia",
            "Permite observar consistencia temporal del sistema. La cola derecha alta afecta el p95 y la experiencia en tiempo real.",
        ),
        (
            "06_scatter_confianza_vs_similitud.png",
            "Confianza detector vs calidad OCR",
            "Mide si detecciones mas confiables tambien producen mejores lecturas OCR. Una tendencia ascendente es deseable.",
        ),
    ]

    for filename, title, text in explanations:
        path = out_dir / "graficas" / filename
        if path.exists():
            doc.add_heading(title, level=2)
            doc.add_paragraph(text)
            doc.add_picture(str(path), width=Inches(6.2))

    doc.add_heading("3. Tablas generadas", level=1)
    doc.add_paragraph("Se generaron las siguientes tablas CSV en la carpeta tablas:")
    doc.add_paragraph("- metricas_detalle_imagen.csv: resultados por imagen (GT, prediccion, latencia, similitud, CER).")
    doc.add_paragraph("- metricas_resumen.csv: indicadores agregados de la corrida actual.")

    latest = (existing.get("validation_metrics_latest") or {}).get("summary", {})
    if latest:
        doc.add_heading("4. Contexto historico", level=1)
        doc.add_paragraph(
            "El proyecto ya tenia validaciones historicas. Se mantienen para trazabilidad y comparacion con la corrida actual."
        )
        doc.add_paragraph(
            f"Historico detection_rate={float(latest.get('detection_rate', 0.0)):.3f}, "
            f"ocr_exact_match_overall={float(latest.get('ocr_exact_match_overall', 0.0)):.3f}, "
            f"mean_detector_confidence={float(latest.get('mean_detector_confidence', 0.0)):.3f}."
        )

    doc.add_heading("5. Conclusiones operativas", level=1)
    doc.add_paragraph(
        "La deteccion se mantiene alta en la corrida evaluada, mientras que el OCR presenta margen de mejora en exactitud total. "
        "Las metricas de latencia deben interpretarse junto al backend OCR activo y al hardware utilizado."
    )

    out_docx = out_dir / "Informe_metricas_y_graficas.docx"
    doc.save(str(out_docx))
    return out_docx


def main():
    repo_root = Path(__file__).resolve().parents[2]
    out_dir = Path(__file__).resolve().parent

    eval_data = run_image_analyzer(repo_root, out_dir, limit=20)
    existing = load_existing_metrics(repo_root)
    plots = generate_plots(out_dir, eval_data, existing)
    report = build_docx(out_dir, eval_data, existing, plots)

    print("OK")
    print(f"Carpeta de resultados: {out_dir}")
    print(f"DOCX: {report}")
    print(f"Resumen JSON: {eval_data['summary_json']}")
    print(f"Tabla detalle: {eval_data['detail_csv']}")


if __name__ == "__main__":
    main()
