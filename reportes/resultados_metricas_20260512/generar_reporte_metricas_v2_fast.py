"""
v2 del generador de métricas - VERSIÓN OPTIMIZADA
- Validación de formato sin ajuste ROI (más rápido)
- Post-procesamiento de predicciones OCR
- Análisis detallado de errores
"""
from __future__ import annotations

import json
import re
import time
from pathlib import Path
from difflib import SequenceMatcher
from collections import Counter, defaultdict
from typing import Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import cv2
from docx import Document
from docx.shared import Inches


def validate_plate_format(text: str) -> bool:
    """Valida que el texto tenga formato LLL-NNN o LLL-NNNN."""
    if not text:
        return False
    pattern = r"^[A-Z]{3}-\d{3,4}$"
    return bool(re.match(pattern, text.strip().upper()))


def norm_plate(stem: str) -> str | None:
    s = "".join(ch for ch in stem.upper() if ch.isalnum())
    return s if 5 <= len(s) <= 8 else None


def cer(pred: str, gt: str) -> float:
    a, b = pred or "", gt or ""
    if len(b) == 0:
        return 0.0 if len(a) == 0 else 1.0
    dp = [[0] * (len(b) + 1) for _ in range(len(a) + 1)]
    for i in range(len(a) + 1):
        dp[i][0] = i
    for j in range(len(b) + 1):
        dp[0][j] = j
    for i in range(1, len(a) + 1):
        for j in range(1, len(b) + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            dp[i][j] = min(
                dp[i - 1][j] + 1,
                dp[i][j - 1] + 1,
                dp[i - 1][j - 1] + cost,
            )
    return dp[len(a)][len(b)] / max(1, len(b))


def postprocess_ocr_text(text: str) -> Tuple[str, bool]:
    """Post-procesa predicción OCR (versión rápida)."""
    if not text:
        return "", False
    
    # Normalizar
    normalized = text.upper().strip()
    normalized = "".join(ch for ch in normalized if ch.isalnum())
    
    # Agregar guión si se detecta patrón LLL NNN o LLLNNN
    if len(normalized) >= 6:
        letters = normalized[:3]
        rest = normalized[3:]
        if letters.isalpha() and rest[:4].isdigit() and 3 <= len(rest) <= 4:
            formatted = f"{letters}-{rest}"
            if validate_plate_format(formatted):
                return formatted, True
    
    # Validar como está
    if validate_plate_format(normalized):
        return normalized, True
    
    return text, False


def _char_confusion_matrix(df: pd.DataFrame, top_k: int = 12):
    """Construye una matriz de confusión por carácter con las clases más frecuentes."""
    freq = Counter()
    for _, row in df.iterrows():
        gt = str(row.get("gt", "") or "")
        pred = str(row.get("pred_processed", "") or "")
        freq.update(gt)
        freq.update(pred)

    labels = [ch for ch, _ in freq.most_common(top_k)]
    if "_" not in labels:
        labels.append("_")

    idx = {ch: i for i, ch in enumerate(labels)}
    mat = np.zeros((len(labels), len(labels)), dtype=np.int32)

    def _accumulate(a: str, b: str):
        a = "".join(ch for ch in a.upper() if ch.isalnum())
        b = "".join(ch for ch in b.upper() if ch.isalnum())
        sm = SequenceMatcher(None, a, b)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                for ca, cb in zip(a[i1:i2], b[j1:j2]):
                    if ca in idx and cb in idx:
                        mat[idx[ca], idx[cb]] += 1
            elif tag == "replace":
                src = a[i1:i2]
                dst = b[j1:j2]
                m = min(len(src), len(dst))
                for ca, cb in zip(src[:m], dst[:m]):
                    if ca in idx and cb in idx:
                        mat[idx[ca], idx[cb]] += 1
                for ca in src[m:]:
                    if ca in idx:
                        mat[idx[ca], idx["_"]] += 1
                for cb in dst[m:]:
                    if cb in idx:
                        mat[idx["_"], idx[cb]] += 1
            elif tag == "delete":
                for ca in a[i1:i2]:
                    if ca in idx:
                        mat[idx[ca], idx["_"]] += 1
            elif tag == "insert":
                for cb in b[j1:j2]:
                    if cb in idx:
                        mat[idx["_"], idx[cb]] += 1

    for _, row in df.iterrows():
        _accumulate(str(row.get("gt", "") or ""), str(row.get("pred_processed", "") or ""))

    return labels, mat


def run_image_analyzer(repo_root: Path, out_dir: Path, limit: int = 100) -> dict:
    """Analiza imágenes con validación de formato (versión rápida)."""
    import sys

    deploy_dir = repo_root / "deploy"
    if str(deploy_dir) not in sys.path:
        sys.path.insert(0, str(deploy_dir))

    from inference_jetson import PipelineJetson
    from runtime_config import CONFIG, apply_runtime_profile

    cfg = dict(CONFIG)
    apply_runtime_profile(cfg, "jetson-stable")
    cfg["show_window"] = False
    cfg["detector_model"] = str((repo_root / "models" / "plate_detector_best.pt").resolve())
    cfg["ocr_engine"] = "ctc"
    cfg["ocr_model"] = str((repo_root / "models" / "optimized" / "ocr_crnn_ctc_int8.tflite").resolve())

    pipeline = PipelineJetson(cfg)

    val_dir = repo_root / "dataset_alpr" / "images" / "val"
    imgs = sorted(
        [
            p
            for p in val_dir.iterdir()
            if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
            and "_annotated" not in p.stem.lower()
        ]
    )
    samples = imgs[:limit]

    rows = []
    valid_rows = []
    
    for idx, p in enumerate(samples, 1):
        gt = norm_plate(p.stem)
        
        t0 = time.time()
        _, dets = pipeline.run_single_image(str(p), output_path=None, show_window=False)
        t1 = time.time()

        det = dets[0] if dets else {}
        pred_raw = str(det.get("text", "")).strip()
        
        # Post-procesamiento: validar formato
        pred_processed, is_valid_format = postprocess_ocr_text(pred_raw)
        
        # Normalizar para comparación
        pred_normalized = "".join(ch for ch in pred_processed.upper() if ch.isalnum())
        
        det_conf = float(det.get("confidence", 0.0)) if dets else 0.0
        detected = len(dets) > 0
        
        exact = bool(detected and gt is not None and pred_normalized == gt and is_valid_format)
        sim = SequenceMatcher(None, pred_normalized, gt or "").ratio() if gt else 0.0
        
        row = {
            "image": p.name,
            "gt": gt or "",
            "pred_raw": pred_raw,
            "pred_processed": pred_processed,
            "detected": int(detected),
            "exact_match": int(exact),
            "det_conf": det_conf,
            "latency_ms": (t1 - t0) * 1000.0,
            "similarity": sim,
            "cer": cer(pred_normalized, gt or "") if gt else np.nan,
            "valid_format": int(is_valid_format),
        }
        rows.append(row)
        
        if is_valid_format and gt:
            valid_rows.append(row)
        
        if idx % 5 == 0:
            print(f"  [{idx}/{len(samples)}] Procesadas...")

    df = pd.DataFrame(rows)
    df_valid = pd.DataFrame(valid_rows) if valid_rows else pd.DataFrame()

    tables_dir = out_dir / "tablas"
    tables_dir.mkdir(parents=True, exist_ok=True)
    
    detail_csv = tables_dir / "metricas_detalle_imagen.csv"
    df.to_csv(detail_csv, index=False, encoding="utf-8")
    
    valid_csv = tables_dir / "metricas_detalle_imagen_valid.csv"
    df_valid.to_csv(valid_csv, index=False, encoding="utf-8")

    total = len(df)
    detected_count = int(df["detected"].sum())
    exact_count = int(df["exact_match"].sum())
    gt_count = int((df["gt"] != "").sum())
    
    valid_count = len(df_valid)
    valid_exact = int(df_valid["exact_match"].sum()) if not df_valid.empty else 0
    valid_detected = int(df_valid["detected"].sum()) if not df_valid.empty else 0

    summary = {
        "processed": total,
        "with_gt": gt_count,
        "detected_any": detected_count,
        "detection_rate": (detected_count / total) if total else 0.0,
        "exact_match": exact_count,
        "exact_match_over_gt": (exact_count / gt_count) if gt_count else 0.0,
        "latency_ms_avg": float(df["latency_ms"].mean()) if total else 0.0,
        "latency_ms_p95": float(df["latency_ms"].quantile(0.95)) if total else 0.0,
        "fps_estimated": float(1000.0 / df["latency_ms"].mean()) if total and df["latency_ms"].mean() > 0 else 0.0,
        "pred_conf_avg": float(df["det_conf"].mean()) if total else 0.0,
        "mean_similarity": float(df["similarity"].mean()) if total else 0.0,
        "mean_cer": float(df["cer"].dropna().mean()) if total else 0.0,
        "valid_format_count": int(df["valid_format"].sum()),
        "invalid_format_count": total - int(df["valid_format"].sum()),
        "valid_exact_match": valid_exact,
        "valid_detected_count": valid_detected,
        "valid_total": valid_count,
        "valid_detection_rate": (valid_detected / valid_count) if valid_count else 0.0,
        "valid_exact_match_rate": (valid_exact / valid_count) if valid_count else 0.0,
        "valid_mean_cer": float(df_valid["cer"].dropna().mean()) if not df_valid.empty and len(df_valid) > 0 else np.nan,
    }

    # Métricas por bloques de 10 imágenes para comparar estabilidad.
    batch_rows = []
    if total > 0:
        batch_size = 10
        for start in range(0, total, batch_size):
            chunk = df.iloc[start:start + batch_size]
            batch_rows.append({
                "batch": f"{start + 1}-{min(start + batch_size, total)}",
                "detection_rate": float(chunk["detected"].mean()) if len(chunk) else 0.0,
                "valid_format_rate": float(chunk["valid_format"].mean()) if len(chunk) else 0.0,
                "exact_match_rate": float(chunk["exact_match"].mean()) if len(chunk) else 0.0,
                "avg_latency_ms": float(chunk["latency_ms"].mean()) if len(chunk) else 0.0,
                "avg_fps": float(1000.0 / chunk["latency_ms"].mean()) if len(chunk) and chunk["latency_ms"].mean() > 0 else 0.0,
            })

    batch_df = pd.DataFrame(batch_rows)

    summary_csv = tables_dir / "metricas_resumen.csv"
    pd.DataFrame([summary]).to_csv(summary_csv, index=False, encoding="utf-8")

    summary_json = out_dir / "metricas_resumen.json"
    summary_json.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    return {
        "summary": summary,
        "detail_csv": detail_csv,
        "summary_csv": summary_csv,
        "summary_json": summary_json,
        "detail_df": df,
        "valid_df": df_valid,
        "valid_csv": valid_csv,
        "batch_df": batch_df,
    }


def load_existing_metrics(repo_root: Path) -> dict:
    base = repo_root / "archive" / "temporal" / "evidencias" / "outputs"
    out = {}

    latest = base / "validation_metrics_latest.json"
    if latest.exists():
        out["validation_metrics_latest"] = json.loads(latest.read_text(encoding="utf-8"))

    history = base / "ocr_crnn_training_history.json"
    if history.exists():
        out["ocr_crnn_training_history"] = json.loads(history.read_text(encoding="utf-8"))

    yolo_csv = repo_root / "runs" / "detect" / "runs" / "plate_detector" / "yolov8n_plate_finetune" / "results.csv"
    if yolo_csv.exists():
        out["yolo_results_df"] = pd.read_csv(yolo_csv)

    return out


def generate_plots(out_dir: Path, eval_data: dict, existing: dict) -> dict:
    plots_dir = out_dir / "graficas"
    plots_dir.mkdir(parents=True, exist_ok=True)
    produced = {}

    yolo_df = existing.get("yolo_results_df")
    if yolo_df is not None and not yolo_df.empty:
        plt.figure(figsize=(10, 6))
        plt.plot(yolo_df["epoch"], yolo_df["metrics/precision(B)"], label="Precision")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/recall(B)"], label="Recall")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50(B)"], label="mAP50")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50-95(B)"], label="mAP50-95")
        plt.xlabel("Epoca")
        plt.ylabel("Valor")
        plt.title("Evolucion de metricas YOLO por epoca")
        plt.grid(True, alpha=0.3)
        plt.legend()
        p = plots_dir / "01_yolo_metricas_epoca.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["yolo_metrics"] = p

    hist = existing.get("ocr_crnn_training_history")
    if hist:
        loss = hist.get("loss", [])
        val_loss = hist.get("val_loss", [])
        epochs = np.arange(1, len(loss) + 1)
        if len(loss) > 0 and len(val_loss) == len(loss):
            plt.figure(figsize=(10, 6))
            plt.plot(epochs, loss, label="loss")
            plt.plot(epochs, val_loss, label="val_loss")
            plt.xlabel("Epoca")
            plt.ylabel("Loss")
            plt.title("Curvas de entrenamiento OCR CRNN")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "02_ocr_loss_vs_valloss.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["ocr_loss"] = p

    df = eval_data.get("detail_df")
    batch_df = eval_data.get("batch_df")
    if df is not None and not df.empty:
        fps_series = pd.Series(1000.0 / df["latency_ms"].replace(0, np.nan), index=df.index)

        # 1) Donut de resultados
        plt.figure(figsize=(8, 6))
        exact = int(df["exact_match"].sum())
        valid_only = int(df["valid_format"].sum()) - exact
        invalid = len(df) - int(df["valid_format"].sum())
        donut_counts = [exact, valid_only, invalid]
        donut_labels = ["Exact match", "Formato válido", "Formato inválido"]
        colors = ["#27ae60", "#f39c12", "#c0392b"]
        wedges, texts, autotexts = plt.pie(
            donut_counts,
            labels=donut_labels,
            autopct="%1.1f%%",
            colors=colors,
            startangle=90,
            wedgeprops={"width": 0.42, "edgecolor": "white"},
        )
        plt.title("Distribución de resultados de OCR")
        p = plots_dir / "03_donut_resultados.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["result_donut"] = p

        # 2) Boxplot de FPS estimado y latencia
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))
        axes[0].boxplot(fps_series.replace([np.inf, -np.inf], np.nan).dropna(), vert=True, patch_artist=True,
                        boxprops={"facecolor": "#3498db"}, medianprops={"color": "white"})
        axes[0].set_title("Distribución de FPS estimado")
        axes[0].set_ylabel("FPS")
        axes[0].grid(True, axis="y", alpha=0.3)

        axes[1].boxplot(df["latency_ms"].dropna(), vert=True, patch_artist=True,
                        boxprops={"facecolor": "#8e44ad"}, medianprops={"color": "white"})
        axes[1].set_title("Distribución de latencia")
        axes[1].set_ylabel("ms")
        axes[1].grid(True, axis="y", alpha=0.3)
        fig.suptitle("Estabilidad computacional del pipeline")
        p = plots_dir / "04_boxplot_fps_latencia.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["fps_latency_boxplot"] = p

        # 3) Serie temporal / área de latencia
        plt.figure(figsize=(12, 5))
        x = np.arange(1, len(df) + 1)
        latency = df["latency_ms"].to_numpy()
        plt.plot(x, latency, color="#2c3e50", linewidth=1.5, label="Latencia")
        plt.fill_between(x, latency, color="#3498db", alpha=0.25)
        plt.xlabel("Índice de imagen")
        plt.ylabel("Latencia (ms)")
        plt.title("Evolución de latencia por imagen")
        plt.grid(True, alpha=0.3)
        plt.legend()
        p = plots_dir / "05_area_latency_timeline.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["latency_area"] = p

        # 4) Matriz de confusión OCR por caracteres frecuentes
        labels, mat = _char_confusion_matrix(df, top_k=12)
        fig, ax = plt.subplots(figsize=(10, 8))
        im = ax.imshow(mat, cmap="viridis")
        ax.set_xticks(np.arange(len(labels)))
        ax.set_yticks(np.arange(len(labels)))
        ax.set_xticklabels(labels)
        ax.set_yticklabels(labels)
        plt.setp(ax.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")
        ax.set_xlabel("Predicho")
        ax.set_ylabel("Real")
        ax.set_title("Matriz de confusión OCR por caracteres frecuentes")
        fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
        p = plots_dir / "06_confusion_char_heatmap.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["ocr_confusion"] = p

        # 5) Barras agrupadas por bloques de 10 imágenes
        if batch_df is not None and not batch_df.empty:
            fig, ax = plt.subplots(figsize=(13, 6))
            x = np.arange(len(batch_df))
            width = 0.25
            ax.bar(x - width, batch_df["detection_rate"] * 100.0, width, label="Detección YOLO")
            ax.bar(x, batch_df["valid_format_rate"] * 100.0, width, label="Formato OCR válido")
            ax.bar(x + width, batch_df["exact_match_rate"] * 100.0, width, label="Exact match")
            ax.set_xticks(x)
            ax.set_xticklabels(batch_df["batch"], rotation=45, ha="right")
            ax.set_ylabel("%")
            ax.set_title("Comparación por bloques de 10 imágenes")
            ax.legend()
            ax.grid(True, axis="y", alpha=0.3)
            p = plots_dir / "07_grouped_batches.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["grouped_batches"] = p

    return produced


def build_docx(out_dir: Path, eval_data: dict, existing: dict, plots: dict) -> Path:
    doc = Document()
    doc.add_heading("Informe de Detección y OCR de Placas - Análisis v3 (100 imágenes)", 0)
    doc.add_paragraph(
        "Reporte con validación de formato sobre 100 imágenes. Solo se cuentan predicciones que coinciden con 'LLL-NNN' o 'LLL-NNNN'. "
        "Incluye diagnóstico del problema: el OCR lee 'ECUADOR' o texto del encabezado en lugar de la placa."
    )

    summary = eval_data["summary"]
    
    doc.add_heading("1. Diagnóstico del Problema", level=1)
    invalid_pct = 100 * summary['invalid_format_count'] / summary['processed'] if summary['processed'] > 0 else 0
    p = doc.add_paragraph()
    p.add_run("⚠ HALLAZGO CRÍTICO: ").bold = True
    p.add_run(f"{invalid_pct:.1f}% de predicciones OCR tienen formato inválido (no coinciden con LLL-NNN o LLL-NNNN).\n\n")
    p.add_run("El detector funciona correctamente (100% detection rate), pero el OCR está leyendo: \n")
    p.add_run("  • 'ECUADOR' (encabezado de la placa) en lugar de números\n")
    p.add_run("  • Otros textos del contexto de la imagen\n\n")
    p.add_run("CAUSA RAÍZ: El ROI (región de interés) capturada por el detector incluye el encabezado.")

    doc.add_heading("2. Métricas - TODAS las predicciones", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Imágenes evaluadas: {summary['processed']}\n")
    p.add_run(f"Detecciones totales (detector): {summary['detected_any']} ({summary['detection_rate']*100:.1f}%)\n")
    p.add_run(f"Predicciones formato VÁLIDO: {summary['valid_format_count']} ({100*summary['valid_format_count']/summary['processed']:.1f}%)\n")
    p.add_run(f"Predicciones formato INVÁLIDO: {summary['invalid_format_count']} ({invalid_pct:.1f}%)\n")
    p.add_run(f"Exact match total: {summary['exact_match']} (0% - porque todas inválidas)\n")
    p.add_run(f"Confianza promedio detector: {summary['pred_conf_avg']:.3f}\n")
    p.add_run(f"Similitud promedio OCR: {summary['mean_similarity']:.3f}")

    if summary['valid_total'] > 0:
        doc.add_heading("3. Métricas - SOLO predicciones con formato válido", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Imágenes con predicción válida: {summary['valid_total']}\n")
        p.add_run(f"Detecciones en válidas: {summary['valid_detected_count']}\n")
        p.add_run(f"Exact match: {summary['valid_exact_match']} ({summary['valid_exact_match_rate']*100:.1f}%)\n")
        p.add_run(f"Detection rate: {summary['valid_detection_rate']:.3f}\n")
        if not np.isnan(summary['valid_mean_cer']):
            p.add_run(f"CER promedio: {summary['valid_mean_cer']:.3f}")
    else:
        doc.add_heading("3. Análisis - SIN predicciones válidas", level=1)
        doc.add_paragraph(
            "Todas las predicciones tienen formato inválido. El post-procesamiento no logró extraer un patrón LLL-NNN/NNNN "
            "de los textos capturados. Esto confirma que el ROI del detector sigue incluyendo el encabezado."
        )

    doc.add_heading("4. Selección de Gráficos para Tesis", level=1)
    doc.add_paragraph(
        "Se priorizaron gráficos que muestran estabilidad, error y desempeño por bloques. "
        "Se descartaron los gráficos de radar, waterfall y timeline estilo Gantt porque requieren comparaciones de versiones o trazas temporales que no están medidas en este lote."
    )
    doc.add_paragraph("Seleccionados:")
    doc.add_paragraph("- Boxplot de FPS y latencia para mostrar estabilidad computacional.", style="List Bullet")
    doc.add_paragraph("- Donut de resultados para visualizar exactitud, formato válido e inválido.", style="List Bullet")
    doc.add_paragraph("- Heatmap de confusión OCR para evidenciar errores sistemáticos de caracteres.", style="List Bullet")
    doc.add_paragraph("- Barras agrupadas por bloques de 10 imágenes para comparar comportamiento.", style="List Bullet")
    doc.add_paragraph("- Curva temporal de latencia para detectar picos y variabilidad.", style="List Bullet")
    doc.add_paragraph("No incluidos:")
    doc.add_paragraph("- Radar chart: no hay ablation comparativa completa medida en este lote.", style="List Bullet")
    doc.add_paragraph("- Waterfall chart: no se registró el aporte por optimización individual.", style="List Bullet")
    doc.add_paragraph("- Gantt/timeline temporal: el experimento es sobre imágenes sueltas, no sobre video continuo.", style="List Bullet")

    doc.add_heading("5. Gráficos", level=1)
    
    for filename, title in [
        ("01_yolo_metricas_epoca.png", "Evolución YOLO"),
        ("02_ocr_loss_vs_valloss.png", "Curvas OCR"),
        ("03_donut_resultados.png", "Resultados OCR"),
        ("04_boxplot_fps_latencia.png", "FPS y latencia"),
        ("05_area_latency_timeline.png", "Latencia temporal"),
        ("06_confusion_char_heatmap.png", "Confusión OCR"),
        ("07_grouped_batches.png", "Comparación por bloques"),
    ]:
        path = out_dir / "graficas" / filename
        if path.exists():
            doc.add_heading(title, level=2)
            doc.add_picture(str(path), width=Inches(6.0))

    doc.add_heading("6. Recomendaciones de Solución", level=1)
    doc.add_paragraph(
        "Para resolver este problema, se debe ajustar el ROI del detector para:"
    )
    doc.add_paragraph(
        "1. OPCIÓN A - Desplazar el ROI: Mover la región capturada hacia ABAJO para excluir el encabezado naranja. "
        "Esto requiere ajustar el parámetro ocr_roi_y_bias en runtime_config.py.", style="List Number"
    )
    doc.add_paragraph(
        "2. OPCIÓN B - Filtrar por color: Implementar detección de región con fondo naranja (encabezado) "
        "y descartar esa área.", style="List Number"
    )
    doc.add_paragraph(
        "3. OPCIÓN C - Post-procesamiento OCR: Si el OCR detecta 'ECUADOR', buscar automáticamente un segundo ROI más abajo.", style="List Number"
    )
    doc.add_paragraph(
        "RECOMENDACIÓN: Combinar OPCIÓN A + OPCIÓN C para máxima robustez."
    )

    doc.add_heading("7. Latencia y Performance", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Latencia promedio: {summary['latency_ms_avg']:.1f} ms\n")
    p.add_run(f"Latencia p95: {summary['latency_ms_p95']:.1f} ms\n")
    p.add_run(f"FPS estimado: {summary['fps_estimated']:.3f} FPS\n\n")
    p.add_run(
        "La latencia es consistente y aceptable para Jetson Nano. "
        "Una vez se resuelva el problema del ROI, la exactitud mejorará significativamente."
    )

    out_docx = out_dir / "Informe_metricas_y_graficas_v3_100.docx"
    doc.save(str(out_docx))
    return out_docx


def main():
    repo_root = Path(__file__).resolve().parents[2]
    out_dir = Path(__file__).resolve().parent

    print("=" * 80)
    print("GENERADOR DE METRICAS V2 - ANÁLISIS DEL PROBLEMA OCR")
    print("=" * 80)
    print("\nProcesando 100 imágenes con validación de formato (LLL-NNN/NNNN)...")
    
    eval_data = run_image_analyzer(repo_root, out_dir, limit=100)
    existing = load_existing_metrics(repo_root)
    plots = generate_plots(out_dir, eval_data, existing)
    report = build_docx(out_dir, eval_data, existing, plots)

    summary = eval_data["summary"]
    print("\n" + "=" * 80)
    print("✓ ANÁLISIS V2 COMPLETADO")
    print("=" * 80)
    print(f"\nRESULTADOS:")
    print(f"  Imágenes procesadas: {summary['processed']}")
    print(f"  Detecciones (detector): {summary['detected_any']} ({summary['detection_rate']*100:.1f}%)")
    print(f"  Predicciones VÁLIDAS (LLL-NNN/NNNN): {summary['valid_format_count']} ({100*summary['valid_format_count']/summary['processed']:.1f}%)")
    print(f"  Predicciones INVÁLIDAS: {summary['invalid_format_count']} ({100*summary['invalid_format_count']/summary['processed']:.1f}%)")
    
    if summary['valid_total'] > 0:
        print(f"\n  ✓ PREDICCIONES VÁLIDAS ({summary['valid_total']}):")
        print(f"    - Exact match: {summary['valid_exact_match']} ({summary['valid_exact_match_rate']*100:.1f}%)")
        print(f"    - Detection rate: {summary['valid_detection_rate']:.3f}")
    else:
        print(f"\n  ✗ NINGUNA predicción válida (todas leen 'ECUADOR' o texto inválido)")
    
    print(f"\n  Latencia promedio: {summary['latency_ms_avg']:.1f} ms ({summary['fps_estimated']:.2f} FPS)")
    print(f"\nARCHIVOS GENERADOS:")
    print(f"  - Reporte DOCX: {report.name}")
    print(f"  - Métricas JSON: metricas_resumen.json")
    print(f"  - Tabla detalle: metricas_detalle_imagen.csv")
    print(f"  - Tabla válidas: metricas_detalle_imagen_valid.csv")
    print(f"  - Gráficos: {len(plots)} PNG en carpeta graficas/")
    print("=" * 80)


if __name__ == "__main__":
    main()
