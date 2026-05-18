"""
v2 del generador de métricas con validación de formato de placa.
- Solo acepta predicciones con formato: LLL-NNN o LLL-NNNN (3 letras - 3/4 números)
- Rechaza predicciones inválidas como "ECUADOR"
- Genera análisis detallado de errores con matriz de confusión
- Reporte mejorado con conclusiones basadas en placas válidas
"""
from __future__ import annotations

import json
import re
import time
from pathlib import Path
from difflib import SequenceMatcher
from collections import Counter

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import cv2
from docx import Document
from docx.shared import Inches


def validate_plate_format(text: str) -> bool:
    """Valida que el texto tenga formato LLL-NNN o LLL-NNNN."""
    if not text:
        return False
    pattern = r"^[A-Z]{3}-\d{3,4}$"
    return bool(re.match(pattern, text.strip().upper()))


def norm_plate(stem: str) -> str | None:
    """Extrae placa normalizada del nombre de archivo."""
    s = "".join(ch for ch in stem.upper() if ch.isalnum())
    return s if 5 <= len(s) <= 8 else None


def cer(pred: str, gt: str) -> float:
    """Calcula Character Error Rate usando distancia de Levenshtein."""
    a, b = pred or "", gt or ""
    if len(b) == 0:
        return 0.0 if len(a) == 0 else 1.0
    dp = [[0] * (len(b) + 1) for _ in range(len(a) + 1)]
    for i in range(len(a) + 1):
        dp[i][0] = i
    for j in range(len(b) + 1):
        dp[0][j] = j
    for i in range(1, len(a) + 1):
        for j in range(1, len(b) + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            dp[i][j] = min(
                dp[i - 1][j] + 1,
                dp[i][j - 1] + 1,
                dp[i - 1][j - 1] + cost,
            )
    return dp[len(a)][len(b)] / max(1, len(b))


def compute_error_breakdown(pred: str, gt: str) -> dict:
    """Desglosa errores en sustituciones, inserciones, eliminaciones."""
    from difflib import SequenceMatcher as SM
    
    errors = {"substitutions": 0, "insertions": 0, "deletions": 0, "char_errors": []}
    
    # Alineación simple de caracteres
    s = SM(None, gt, pred)
    errors["char_errors"] = [
        {"gt_char": g, "pred_char": p, "type": "substitution"} 
        for g, p in zip(gt, pred[:len(gt)])
        if g != p
    ]
    
    if len(pred) > len(gt):
        errors["insertions"] = len(pred) - len(gt)
    elif len(gt) > len(pred):
        errors["deletions"] = len(gt) - len(pred)
    
    return errors


def build_confusion_matrix(detail_df: pd.DataFrame) -> dict:
    """Construye matriz de confusión de caracteres."""
    confusion = Counter()
    
    for _, row in detail_df.iterrows():
        if pd.isna(row.get("cer")) or row.get("cer") is None:
            continue
        gt = str(row.get("gt", ""))
        pred = str(row.get("pred", ""))
        
        # Mapear caracteres mal reconocidos
        for i, (g, p) in enumerate(zip(gt, pred[:len(gt)])):
            if g != p:
                confusion[(g, p)] += 1
    
    return dict(confusion)


def run_image_analyzer(repo_root: Path, out_dir: Path, limit: int = 20) -> dict:
    """Analiza imágenes de validación con validación de formato de placa y ajuste ROI."""
    import sys

    deploy_dir = repo_root / "deploy"
    if str(deploy_dir) not in sys.path:
        sys.path.insert(0, str(deploy_dir))

    from inference_jetson import PipelineJetson
    from runtime_config import CONFIG, apply_runtime_profile
    from ocr_postprocessor import (
        postprocess_ocr_prediction,
        normalize_plate_text,
        extract_plate_from_text,
        validate_plate_format as validate_format
    )

    cfg = dict(CONFIG)
    apply_runtime_profile(cfg, "jetson-stable")
    cfg["show_window"] = False
    cfg["detector_model"] = str((repo_root / "models" / "plate_detector_best.pt").resolve())
    cfg["ocr_engine"] = "ctc"
    cfg["ocr_model"] = str((repo_root / "models" / "optimized" / "ocr_crnn_ctc_int8.tflite").resolve())

    pipeline = PipelineJetson(cfg)

    val_dir = repo_root / "dataset_alpr" / "images" / "val"
    ann_dir = out_dir / "anotadas"
    ann_dir.mkdir(parents=True, exist_ok=True)
    imgs = sorted(
        [
            p
            for p in val_dir.iterdir()
            if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
            and "_annotated" not in p.stem.lower()
        ]
    )
    samples = imgs[:limit]

    rows = []
    valid_rows = []  # Separadas por formato válido
    
    for p in samples:
        gt = norm_plate(p.stem)
        frame = cv2.imread(str(p))
        
        t0 = time.time()
        ann_path = ann_dir / f"{p.stem}_annotated{p.suffix}"
        _, dets = pipeline.run_single_image(str(p), output_path=str(ann_path), show_window=False)
        t1 = time.time()

        det = dets[0] if dets else {}
        pred_raw = str(det.get("text", "")).strip()
        
        # POST-PROCESAMIENTO: Ajuste ROI si OCR leyó "ECUADOR"
        pred_post, is_valid_format, _ = postprocess_ocr_prediction(pred_raw)
        
        # Si la predicción sigue siendo inválida pero contiene "ECUADOR", 
        # intentar ajustar el ROI manualmente
        if not is_valid_format and "ECUADOR" in pred_raw.upper() and frame is not None:
            # Estrategia: Bajar el ROI
            bbox = det.get("bbox", (0, 0, 100, 100))
            if len(bbox) >= 4:
                x1, y1, x2, y2 = bbox[:4]
                roi_height = y2 - y1
                
                # Intentar predicción con ROI desplazado hacia abajo
                for offset_pct in [0.25, 0.35, 0.15]:
                    y_offset = int(roi_height * offset_pct)
                    y1_new = min(y1 + y_offset, frame.shape[0] - 1)
                    y2_new = min(y2 + y_offset, frame.shape[0])
                    
                    if y2_new <= y1_new:
                        continue
                    
                    try:
                        roi_adjusted = frame[y1_new:y2_new, x1:x2]
                        pred_temp_raw, _ = pipeline.ocr.segment_and_recognize_with_conf(roi_adjusted)
                        pred_temp, is_valid_temp, _ = postprocess_ocr_prediction(pred_temp_raw)
                        
                        if is_valid_temp:
                            pred_post = pred_temp
                            is_valid_format = True
                            break
                    except Exception:
                        pass
        
        pred = pred_post
        pred_normalized = normalize_plate_text(pred)
        
        det_conf = float(det.get("confidence", 0.0)) if dets else 0.0
        detected = len(dets) > 0
        
        exact = bool(detected and gt is not None and pred_normalized == gt and is_valid_format)
        sim = SequenceMatcher(None, pred_normalized, gt or "").ratio() if gt else 0.0
        
        row = {
            "image": p.name,
            "gt": gt or "",
            "pred": pred_normalized,
            "pred_raw": pred_raw,
            "pred_post": pred,
            "detected": int(detected),
            "exact_match": int(exact),
            "det_conf": det_conf,
            "latency_ms": (t1 - t0) * 1000.0,
            "similarity": sim,
            "cer": cer(pred_normalized, gt or "") if gt else np.nan,
            "valid_format": int(is_valid_format),
        }
        rows.append(row)
        
        # Separar solo placas con formato válido
        if is_valid_format and gt:
            valid_rows.append(row)

    df = pd.DataFrame(rows)
    df_valid = pd.DataFrame(valid_rows) if valid_rows else pd.DataFrame()

    tables_dir = out_dir / "tablas"
    tables_dir.mkdir(parents=True, exist_ok=True)
    
    detail_csv = tables_dir / "metricas_detalle_imagen.csv"
    df.to_csv(detail_csv, index=False, encoding="utf-8")
    
    valid_csv = tables_dir / "metricas_detalle_imagen_valid.csv"
    df_valid.to_csv(valid_csv, index=False, encoding="utf-8")

    total = len(df)
    detected_count = int(df["detected"].sum())
    exact_count = int(df["exact_match"].sum())
    gt_count = int((df["gt"] != "").sum())
    
    valid_count = len(df_valid)
    valid_exact = int(df_valid["exact_match"].sum()) if not df_valid.empty else 0
    valid_detected = int(df_valid["detected"].sum()) if not df_valid.empty else 0

    summary = {
        "processed": total,
        "with_gt": gt_count,
        "detected_any": detected_count,
        "detection_rate": (detected_count / total) if total else 0.0,
        "exact_match": exact_count,
        "exact_match_over_gt": (exact_count / gt_count) if gt_count else 0.0,
        "latency_ms_avg": float(df["latency_ms"].mean()) if total else 0.0,
        "latency_ms_p95": float(df["latency_ms"].quantile(0.95)) if total else 0.0,
        "fps_estimated": float(1000.0 / df["latency_ms"].mean()) if total and df["latency_ms"].mean() > 0 else 0.0,
        "pred_conf_avg": float(df["det_conf"].mean()) if total else 0.0,
        "mean_similarity": float(df["similarity"].mean()) if total else 0.0,
        "mean_cer": float(df["cer"].dropna().mean()) if total else 0.0,
        "valid_format_count": int(df["valid_format"].sum()),
        "invalid_format_count": total - int(df["valid_format"].sum()),
        "valid_exact_match": valid_exact,
        "valid_detected_count": valid_detected,
        "valid_total": valid_count,
        "valid_detection_rate": (valid_detected / valid_count) if valid_count else 0.0,
        "valid_exact_match_rate": (valid_exact / valid_count) if valid_count else 0.0,
        "valid_mean_cer": float(df_valid["cer"].dropna().mean()) if not df_valid.empty and len(df_valid) > 0 else np.nan,
    }

    summary_csv = tables_dir / "metricas_resumen.csv"
    pd.DataFrame([summary]).to_csv(summary_csv, index=False, encoding="utf-8")

    summary_json = out_dir / "metricas_resumen.json"
    summary_json.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    return {
        "summary": summary,
        "detail_csv": detail_csv,
        "summary_csv": summary_csv,
        "summary_json": summary_json,
        "detail_df": df,
        "valid_df": df_valid,
        "valid_csv": valid_csv,
    }


def load_existing_metrics(repo_root: Path) -> dict:
    """Carga métricas históricas."""
    base = repo_root / "archive" / "temporal" / "evidencias" / "outputs"
    out = {}

    latest = base / "validation_metrics_latest.json"
    if latest.exists():
        out["validation_metrics_latest"] = json.loads(latest.read_text(encoding="utf-8"))

    history = base / "ocr_crnn_training_history.json"
    if history.exists():
        out["ocr_crnn_training_history"] = json.loads(history.read_text(encoding="utf-8"))

    yolo_csv = repo_root / "runs" / "detect" / "runs" / "plate_detector" / "yolov8n_plate_finetune" / "results.csv"
    if yolo_csv.exists():
        out["yolo_results_df"] = pd.read_csv(yolo_csv)

    return out


def generate_plots(out_dir: Path, eval_data: dict, existing: dict) -> dict:
    """Genera gráficos incluyendo análisis de formato válido."""
    plots_dir = out_dir / "graficas"
    plots_dir.mkdir(parents=True, exist_ok=True)

    produced = {}

    # 1) YOLO metrics by epoch
    yolo_df = existing.get("yolo_results_df")
    if yolo_df is not None and not yolo_df.empty:
        plt.figure(figsize=(10, 6))
        plt.plot(yolo_df["epoch"], yolo_df["metrics/precision(B)"], label="Precision")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/recall(B)"], label="Recall")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50(B)"], label="mAP50")
        plt.plot(yolo_df["epoch"], yolo_df["metrics/mAP50-95(B)"], label="mAP50-95")
        plt.xlabel("Epoca")
        plt.ylabel("Valor")
        plt.title("Evolucion de metricas YOLO por epoca")
        plt.grid(True, alpha=0.3)
        plt.legend()
        p = plots_dir / "01_yolo_metricas_epoca.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["yolo_metrics"] = p

    # 2) OCR training curves
    hist = existing.get("ocr_crnn_training_history")
    if hist:
        loss = hist.get("loss", [])
        val_loss = hist.get("val_loss", [])
        epochs = np.arange(1, len(loss) + 1)
        if len(loss) > 0 and len(val_loss) == len(loss):
            plt.figure(figsize=(10, 6))
            plt.plot(epochs, loss, label="loss")
            plt.plot(epochs, val_loss, label="val_loss")
            plt.xlabel("Epoca")
            plt.ylabel("Loss")
            plt.title("Curvas de entrenamiento OCR CRNN")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "02_ocr_loss_vs_valloss.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["ocr_loss"] = p

        blank = hist.get("val_blank_rate", [])
        empty = hist.get("val_empty_pred_rate", [])
        if len(blank) == len(empty) and len(blank) > 0:
            plt.figure(figsize=(10, 6))
            plt.plot(np.arange(1, len(blank) + 1), blank, label="val_blank_rate")
            plt.plot(np.arange(1, len(empty) + 1), empty, label="val_empty_pred_rate")
            plt.xlabel("Epoca")
            plt.ylabel("Tasa")
            plt.title("Control de colapso CTC en validacion")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "03_ocr_blank_empty_rate.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["ocr_collapse"] = p

    # 3) Comparativo antes y después de filtro de formato
    df = eval_data.get("detail_df")
    if df is not None and not df.empty:
        categories = ["Detection rate", "Exact match", "Confidence"]
        all_metrics = [
            df["detected"].mean(),
            (df["exact_match"].sum() / len(df)) if len(df) > 0 else 0,
            df["det_conf"].mean(),
        ]
        valid_metrics = []
        valid_df = eval_data.get("valid_df")
        if valid_df is not None and len(valid_df) > 0:
            valid_metrics = [
                valid_df["detected"].mean(),
                (valid_df["exact_match"].sum() / len(valid_df)) if len(valid_df) > 0 else 0,
                valid_df["det_conf"].mean(),
            ]
        
        if valid_metrics:
            x = np.arange(len(categories))
            w = 0.35
            plt.figure(figsize=(10, 6))
            plt.bar(x - w / 2, all_metrics, width=w, label="Todas las predicciones")
            plt.bar(x + w / 2, valid_metrics, width=w, label="Solo formato valido (LLL-NNN/NNNN)")
            plt.xticks(x, categories)
            plt.ylim(0, 1.05)
            plt.ylabel("Valor")
            plt.title("Impacto del filtro de formato valido")
            plt.grid(True, axis="y", alpha=0.3)
            plt.legend()
            p = plots_dir / "04_comparativo_antes_despues_filtro.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["format_filter_comparison"] = p

        # 4) Latency distribution
        plt.figure(figsize=(10, 6))
        plt.hist(df["latency_ms"], bins=10)
        plt.xlabel("Latencia por imagen (ms)")
        plt.ylabel("Frecuencia")
        plt.title("Distribucion de latencia del analizador de imagenes")
        plt.grid(True, alpha=0.3)
        p = plots_dir / "05_histograma_latencia.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["latency_hist"] = p

        # 5) Confianza vs similitud (solo válidas)
        if valid_df is not None and len(valid_df) > 0:
            plt.figure(figsize=(10, 6))
            plt.scatter(valid_df["det_conf"], valid_df["similarity"], alpha=0.8, color="green", label="Formato válido")
            invalid_df = df[df["valid_format"] == 0]
            if len(invalid_df) > 0:
                plt.scatter(invalid_df["det_conf"], invalid_df["similarity"], alpha=0.4, color="red", label="Formato inválido")
            plt.xlabel("Confianza detector")
            plt.ylabel("Similitud OCR vs GT")
            plt.title("Confianza detector vs calidad OCR (filtrado por formato)")
            plt.grid(True, alpha=0.3)
            plt.legend()
            p = plots_dir / "06_scatter_confianza_vs_similitud_filtrado.png"
            plt.tight_layout()
            plt.savefig(p, dpi=150)
            plt.close()
            produced["conf_vs_similarity"] = p

        # 6) Gráfico de formato válido vs inválido
        plt.figure(figsize=(8, 6))
        format_counts = [int(df["valid_format"].sum()), len(df) - int(df["valid_format"].sum())]
        plt.pie(format_counts, labels=["Formato válido", "Formato inválido"], autopct="%1.1f%%", colors=["green", "red"])
        plt.title("Distribución de predicciones por validez de formato")
        p = plots_dir / "07_distribucion_formato.png"
        plt.tight_layout()
        plt.savefig(p, dpi=150)
        plt.close()
        produced["format_distribution"] = p

    return produced


def build_docx(out_dir: Path, eval_data: dict, existing: dict, plots: dict) -> Path:
    """Construye reporte DOCX con análisis de formato validado."""
    doc = Document()
    doc.add_heading("Informe de Resultados - Deteccion y OCR de Placas (v2)", 0)
    doc.add_paragraph(
        "Este documento resume las metricas cuantitativas del proyecto con VALIDACION DE FORMATO. "
        "Solo se consideran las predicciones OCR que coinciden con el patrón 'LLL-NNN' o 'LLL-NNNN'."
    )

    summary = eval_data["summary"]
    
    # Resumen General
    doc.add_heading("1. Resumen de metricas - TODAS LAS PREDICCIONES", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Imagenes evaluadas: {summary['processed']}\n")
    p.add_run(f"Detecciones totales: {summary['detected_any']}\n")
    p.add_run(f"Predicciones con formato válido: {summary['valid_format_count']} ({100*summary['valid_format_count']/summary['processed']:.1f}%)\n")
    p.add_run(f"Predicciones con formato inválido: {summary['invalid_format_count']} ({100*summary['invalid_format_count']/summary['processed']:.1f}%)\n")
    p.add_run(f"Detection rate (detector): {summary['detection_rate']:.3f}\n")
    p.add_run(f"Latencia promedio: {summary['latency_ms_avg']:.1f} ms\n")
    p.add_run(f"FPS estimado: {summary['fps_estimated']:.3f}\n")

    # Resumen Filtrado
    doc.add_heading("2. Metricas FILTRADAS - Solo formato válido (LLL-NNN/NNNN)", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Imagenes con formato válido: {summary['valid_total']}\n")
    p.add_run(f"Detecciones en válidas: {summary['valid_detected_count']}\n")
    p.add_run(f"Exact match OCR (válidas): {summary['valid_exact_match']} ({summary['valid_exact_match_rate']*100:.1f}%)\n")
    p.add_run(f"Detection rate (válidas): {summary['valid_detection_rate']:.3f}\n")
    cer_text = f"{summary['valid_mean_cer']:.3f}" if not np.isnan(summary['valid_mean_cer']) else "N/A"
    p.add_run(f"CER promedio (válidas): {cer_text}\n")
    p.add_run(f"Confianza promedio detector: {summary['pred_conf_avg']:.3f}\n")
    p.add_run(f"Similitud promedio OCR: {summary['mean_similarity']:.3f}")

    # Gráficos
    doc.add_heading("3. Graficas y analisis", level=1)

    explanations = [
        (
            "01_yolo_metricas_epoca.png",
            "Evolucion de metricas YOLO",
            "Muestra la convergencia del detector. Debe observarse estabilidad en mAP50/mAP50-95.",
        ),
        (
            "02_ocr_loss_vs_valloss.png",
            "Curvas de entrenamiento OCR",
            "Permite verificar aprendizaje y posible sobreajuste del modelo OCR CRNN.",
        ),
        (
            "03_ocr_blank_empty_rate.png",
            "Tasas de blank y prediccion vacia",
            "Evalúan colapso CTC. Una reducción indica que el OCR evita secuencias vacias.",
        ),
        (
            "04_comparativo_antes_despues_filtro.png",
            "Impacto del filtro de formato",
            "Compara métricas ANTES y DESPUÉS de aplicar validación de formato LLL-NNN/NNNN. "
            "Muestra cómo el filtro mejora exactitud al rechazar predicciones inválidas.",
        ),
        (
            "05_histograma_latencia.png",
            "Distribucion de latencia",
            "Permite observar consistencia temporal. La cola derecha afecta el p95 y experiencia real.",
        ),
        (
            "06_scatter_confianza_vs_similitud_filtrado.png",
            "Confianza detector vs calidad OCR (filtrado)",
            "Verde: formato válido. Rojo: inválido. Muestra si detecciones confiables producen mejores OCR.",
        ),
        (
            "07_distribucion_formato.png",
            "Distribución por validez de formato",
            "Porcentaje de predicciones con formato válido vs inválido. Objetivo: maximizar válidas.",
        ),
    ]

    for filename, title, text in explanations:
        path = out_dir / "graficas" / filename
        if path.exists():
            doc.add_heading(title, level=2)
            doc.add_paragraph(text)
            doc.add_picture(str(path), width=Inches(6.2))

    # Análisis detallado
    doc.add_heading("4. Analisis detallado de errores", level=1)
    valid_df = eval_data.get("valid_df")
    if valid_df is not None and len(valid_df) > 0:
        doc.add_paragraph(
            f"De las {len(valid_df)} predicciones con formato válido, {summary['valid_exact_match']} "
            f"coinciden exactamente con el ground truth ({summary['valid_exact_match_rate']*100:.1f}%)."
        )
        
        # Errores por tipo
        errors_by_type = {"substitutions": 0, "insertions": 0, "deletions": 0}
        for _, row in valid_df.iterrows():
            if not pd.isna(row.get("cer")):
                gt = str(row.get("gt", ""))
                pred = str(row.get("pred", ""))
                if len(pred) > len(gt):
                    errors_by_type["insertions"] += len(pred) - len(gt)
                elif len(gt) > len(pred):
                    errors_by_type["deletions"] += len(gt) - len(pred)
                for g, p in zip(gt, pred[:len(gt)]):
                    if g != p:
                        errors_by_type["substitutions"] += 1
        
        p = doc.add_paragraph()
        p.add_run("Desglose de errores OCR:\n")
        p.add_run(f"  - Sustituciones de caracteres: {errors_by_type['substitutions']}\n")
        p.add_run(f"  - Inserciones: {errors_by_type['insertions']}\n")
        p.add_run(f"  - Eliminaciones: {errors_by_type['deletions']}\n")
    
    # Tablas generadas
    doc.add_heading("5. Tablas generadas", level=1)
    doc.add_paragraph("Se generaron las siguientes tablas CSV:")
    doc.add_paragraph("- metricas_detalle_imagen.csv: Todos los resultados (incluye valid_format, pred_raw)")
    doc.add_paragraph("- metricas_detalle_imagen_valid.csv: Solo predicciones con formato válido")
    doc.add_paragraph("- metricas_resumen.csv: Indicadores agregados (incluye métricas filtradas)")

    # Conclusiones
    doc.add_heading("6. Conclusiones operativas", level=1)
    invalid_pct = 100 * summary['invalid_format_count'] / summary['processed'] if summary['processed'] > 0 else 0
    doc.add_paragraph(
        f"HALLAZGO CLAVE: {invalid_pct:.1f}% de predicciones tienen formato inválido "
        f"(texto arbitrario como 'ECUADOR'). Solo {100 - invalid_pct:.1f}% tienen patrón LLL-NNN/NNNN."
    )
    
    if summary['valid_exact_match_rate'] > 0.5:
        doc.add_paragraph(
            f"✓ El OCR con formato válido muestra {summary['valid_exact_match_rate']*100:.1f}% exactitud. "
            "El detector funciona correctamente pero hay falsos positivos en zonas de texto."
        )
    else:
        doc.add_paragraph(
            f"⚠ Incluso con formato válido, exactitud es {summary['valid_exact_match_rate']*100:.1f}%. "
            "Se recomienda: (1) revisar ROI del detector, (2) mejorar OCR en caracteres confundibles, "
            "(3) aumentar datos de entrenamiento para casos edge."
        )
    
    doc.add_paragraph(
        f"Latencia promedio: {summary['latency_ms_avg']:.1f} ms ({summary['fps_estimated']:.2f} FPS). "
        "Aceptable para producción si hardware es Jetson o GPU."
    )

    out_docx = out_dir / "Informe_metricas_y_graficas_v2.docx"
    doc.save(str(out_docx))
    return out_docx


def main():
    repo_root = Path(__file__).resolve().parents[2]
    out_dir = Path(__file__).resolve().parent

    print("=" * 80)
    print("GENERADOR DE METRICAS V2 - CON AJUSTE DE ROI Y POST-PROCESAMIENTO")
    print("=" * 80)
    
    eval_data = run_image_analyzer(repo_root, out_dir, limit=20)
    existing = load_existing_metrics(repo_root)
    plots = generate_plots(out_dir, eval_data, existing)
    report = build_docx(out_dir, eval_data, existing, plots)

    summary = eval_data["summary"]
    print("\n✓ v2 completada con validación de formato + ajuste ROI")
    print(f"\nRESULTADOS:")
    print(f"  Imágenes procesadas: {summary['processed']}")
    print(f"  Detecciones: {summary['detected_any']} ({summary['detection_rate']*100:.1f}%)")
    print(f"  Predicciones formato válido (LLL-NNN/NNNN): {summary['valid_format_count']} ({100*summary['valid_format_count']/summary['processed']:.1f}%)")
    print(f"  Predicciones formato inválido: {summary['invalid_format_count']} ({100*summary['invalid_format_count']/summary['processed']:.1f}%)")
    if summary['valid_total'] > 0:
        print(f"\n  SOLO VÁLIDAS ({summary['valid_total']} imágenes):")
        print(f"    Exact match: {summary['valid_exact_match']} ({summary['valid_exact_match_rate']*100:.1f}%)")
        print(f"    Detection rate: {summary['valid_detection_rate']:.3f}")
        if not np.isnan(summary['valid_mean_cer']):
            print(f"    Mean CER: {summary['valid_mean_cer']:.3f}")
    
    print(f"\nLatencia promedio: {summary['latency_ms_avg']:.1f} ms ({summary['fps_estimated']:.2f} FPS)")
    print(f"\nCarpeta de resultados: {out_dir}")
    print(f"DOCX v2: {report}")
    print(f"Resumen JSON: {eval_data['summary_json']}")
    print(f"Tabla detalle (todas): {eval_data['detail_csv']}")
    print(f"Tabla detalle (válidas): {eval_data['valid_csv']}")
    print("=" * 80)


if __name__ == "__main__":
    main()
